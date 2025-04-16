import sys
import os
import datetime
import keyboard
import pyautogui
from PIL import Image
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
from PyQt5.QtWidgets import (
    QApplication, QWidget, QLabel, QLineEdit, QPushButton, QTextEdit,
    QVBoxLayout, QHBoxLayout, QFileDialog, QCheckBox, QSpinBox, QGroupBox,
    QMessageBox, QComboBox, QListWidget, QListWidgetItem
)
from PyQt5.QtGui import QPixmap, QScreen
from PyQt5.QtCore import Qt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import openpyxl
from openpyxl.drawing.image import Image as ExcelImage
import logging  # Import the logging module


class ScreenshotApp(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Screenshot Tool")
        self.screenshot_count = 1
        self.capture_enabled = False
        self.default_key = 'home'
        self.doc_path = None
        self.doc = None
        self.captured_data = []  # To store data for Excel
        self.captured_images = []  # To keep track of captured image paths
        self.delete_images_after_save = False
        self.selected_monitors = []  # List of selected monitor indices
        self.capture_mode = "single"  # "single", "all", "multiple"

        # Configure logging
        logging.basicConfig(filename='screenshot_app.log', level=logging.DEBUG,
                            format='%(asctime)s - %(levelname)s - %(message)s')
        logging.info("Application started")  # Log application startup

        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout()

        # Group: Basic Settings
        basic_group = QGroupBox("Settings")
        basic_layout = QVBoxLayout()

        basic_layout.addWidget(QLabel("Test Case Name:"))
        self.test_case_input = QLineEdit()
        self.test_case_input.setPlaceholderText("Enter Case Name (Default: Evidence)")
        basic_layout.addWidget(self.test_case_input)

        basic_layout.addWidget(QLabel("Document Version (e.g., v1):"))
        self.version_input = QLineEdit("v1")
        basic_layout.addWidget(self.version_input)

        basic_layout.addWidget(QLabel("Screenshot Key:"))
        self.hotkey_input = QLineEdit(self.default_key)
        basic_layout.addWidget(self.hotkey_input)

        basic_group.setLayout(basic_layout)
        layout.addWidget(basic_group)

        # Group: Monitor Selection
        monitor_group = QGroupBox("Monitor Selection")
        monitor_layout = QVBoxLayout()

        self.monitor_mode_combo = QComboBox()
        self.monitor_mode_combo.addItem("Single Monitor")
        self.monitor_mode_combo.addItem("Capture All Monitors (Stitched)")
        self.monitor_mode_combo.addItem("Select Multiple Monitors")
        self.monitor_mode_combo.currentIndexChanged.connect(self.monitor_mode_changed)
        monitor_layout.addWidget(self.monitor_mode_combo)

        self.single_monitor_combo = QComboBox()
        self.populate_single_monitor_combo()
        monitor_layout.addWidget(self.single_monitor_combo)

        self.multiple_monitor_list = QListWidget()
        self.populate_multiple_monitor_list()
        monitor_layout.addWidget(self.multiple_monitor_list)

        self.update_monitor_visibility()
        monitor_group.setLayout(monitor_layout)
        layout.addWidget(monitor_group)

        # Group: Output
        output_group = QGroupBox("Output Settings")
        output_layout = QVBoxLayout()

        output_layout.addWidget(QLabel("Output Folder:"))
        folder_layout = QHBoxLayout()
        self.folder_input = QLineEdit("/path/to/output/folder")
        browse_button = QPushButton("Browse...")
        browse_button.clicked.connect(self.browse_folder)
        folder_layout.addWidget(self.folder_input)
        folder_layout.addWidget(browse_button)
        output_layout.addLayout(folder_layout)

        self.timestamp_checkbox = QCheckBox("Add Timestamp to Description")
        self.increment_checkbox = QCheckBox("Enable Auto-increment Co.")
        self.increment_spin = QSpinBox()
        self.increment_spin.setValue(1)
        self.delete_checkbox = QCheckBox("Delete Images After Save")
        self.generate_excel_checkbox = QCheckBox("Generate Excel Document with Images")

        output_layout.addWidget(self.timestamp_checkbox)
        output_layout.addWidget(self.increment_checkbox)
        inc_layout = QHBoxLayout()
        inc_layout.addWidget(QLabel("Increment By:"))
        inc_layout.addWidget(self.increment_spin)
        output_layout.addLayout(inc_layout)
        output_layout.addWidget(self.delete_checkbox)
        output_layout.addWidget(self.generate_excel_checkbox)

        output_group.setLayout(output_layout)
        layout.addWidget(output_group)

        # Group: Screenshot Description
        layout.addWidget(QLabel("Screenshot Description:"))
        self.description_input = QLineEdit()
        layout.addWidget(self.description_input)

        # Group: Screenshot Preview
        layout.addWidget(QLabel("Screenshot Preview:"))
        self.preview_label = QLabel()
        self.preview_label.setFixedSize(300, 200)
        self.preview_label.setStyleSheet("border: 1px solid black;")
        layout.addWidget(self.preview_label)

        # Buttons
        button_layout = QHBoxLayout()
        self.start_button = QPushButton("Start New Capture")
        self.append_button = QPushButton("Append to Existing")
        self.stop_button = QPushButton("End Capture & Save")
        self.convert_pdf_button = QPushButton("Convert to PDF")
        self.start_button.clicked.connect(self.start_new_capture)
        self.append_button.clicked.connect(self.append_to_existing)
        self.stop_button.clicked.connect(self.stop_capture)
        self.convert_pdf_button.clicked.connect(self.convert_to_pdf)
        button_layout.addWidget(self.start_button)
        button_layout.addWidget(self.append_button)
        button_layout.addWidget(self.stop_button)
        button_layout.addWidget(self.convert_pdf_button)
        layout.addLayout(button_layout)

        # Status
        self.status_label = QLabel("Status: Ready")
        layout.addWidget(self.status_label)

        self.setLayout(layout)

    def populate_single_monitor_combo(self):
        self.single_monitor_combo.clear()
        screens = QApplication.screens()
        for i, screen in enumerate(screens):
            self.single_monitor_combo.addItem(f"Monitor {i + 1}")
        logging.info("Populated single monitor combo box")

    def populate_multiple_monitor_list(self):
        self.multiple_monitor_list.clear()
        screens = QApplication.screens()
        for i, screen in enumerate(screens):
            item = QListWidgetItem(f"Monitor {i + 1}")
            item.setFlags(item.flags() | Qt.ItemIsUserCheckable)
            item.setCheckState(Qt.Unchecked)
            self.multiple_monitor_list.addItem(item)
        logging.info("Populated multiple monitor list")

    def update_monitor_visibility(self):
        mode = self.monitor_mode_combo.currentText()
        self.single_monitor_combo.setVisible(mode == "Single Monitor")
        self.multiple_monitor_list.setVisible(mode == "Select Multiple Monitors")
        logging.info(f"Updated monitor visibility. Mode: {mode}")

    def monitor_mode_changed(self, index):
        self.capture_mode = ["single", "all", "multiple"][index]
        self.update_monitor_visibility()
        logging.info(f"Monitor mode changed to: {self.capture_mode}")

    def browse_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "Select Output Folder")
        if folder:
            self.folder_input.setText(folder)
            logging.info(f"Output folder selected: {folder}")

    def start_new_capture(self):
        try:
            self.capture_enabled = True
            self.screenshot_count = 1
            self.status_label.setText("Status: Capture Started (New Document)")
            self.hotkey = self.hotkey_input.text().strip().lower()
            keyboard.add_hotkey(self.hotkey, self.capture_screenshot)

            folder = self.folder_input.text().strip()
            case_name = self.test_case_input.text().strip() or "Evidence"
            version = self.version_input.text().strip() or "v1"
            self.doc_path = os.path.join(folder, f"{case_name}_{version}.docx")
            self.excel_path = os.path.join(folder, f"{case_name}_{version}.xlsx")
            self.doc = Document()
            self.captured_data = []
            self.captured_images = []
            self.delete_images_after_save = self.delete_checkbox.isChecked()

            logging.info(f"Started new capture. Document path: {self.doc_path}, Excel path: {self.excel_path}")
        except Exception as e:
            logging.error(f"Error starting new capture: {e}")
            QMessageBox.critical(self, "Error", f"Error starting new capture: {e}")

    def append_to_existing(self):
        try:
            options = QFileDialog.Options()
            options |= QFileDialog.DontUseNativeDialog
            file_path, _ = QFileDialog.getOpenFileName(self, "Select Existing Word Document", "",
                                                      "Word Documents (*.docx)", options=options)
            if file_path:
                self.doc_path = file_path
                self.doc = Document(self.doc_path)
                self.capture_enabled = True
                self.status_label.setText(
                    f"Status: Capture Started (Appending to {os.path.basename(self.doc_path)})")
                self.hotkey = self.hotkey_input.text().strip().lower()
                keyboard.add_hotkey(self.hotkey, self.capture_screenshot)
                last_paragraph = self.doc.paragraphs[-1].text if self.doc.paragraphs else ""
                if "Screenshot" in last_paragraph:
                    try:
                        self.screenshot_count = int(last_paragraph.split("Screenshot ")[1].split(":")[0]) + 1
                    except ValueError:
                        self.screenshot_count = 1
                else:
                    self.screenshot_count = 1
                self.captured_data = []  # Reset Excel data for append session
                self.captured_images = []  # Reset image list for this append session
                self.delete_images_after_save = self.delete_checkbox.isChecked()
                logging.info(f"Appending to existing document: {self.doc_path}")
        except Exception as e:
            logging.error(f"Error appending to existing document: {e}")
            QMessageBox.critical(self, "Error", f"Could not open Word document: {e}")
            self.doc_path = None
            self.doc = None

    def stop_capture(self):
        try:
            self.capture_enabled = False
            keyboard.unhook_all_hotkeys()

            if self.doc and self.doc_path:
                self.doc.save(self.doc_path)
                self.status_label.setText(f"Capture complete. Word document saved to: {self.doc_path}")
                logging.info(f"Capture complete. Word document saved to: {self.doc_path}")
                if self.generate_excel_checkbox.isChecked() and self.captured_data:
                    self.generate_excel()

                # Convert to PDF (optional, only if a new document was started)
                if "New Document" in self.status_label.text():
                    try:
                        pdf_path = os.path.splitext(self.doc_path)[0] + ".pdf"
                        convert(self.doc_path, pdf_path)
                        self.status_label.setText(f"Word and PDF saved.")
                        logging.info(f"Word and PDF saved. PDF path: {pdf_path}")
                    except Exception as e:
                        logging.error(f"PDF conversion failed: {e}")
                        self.status_label.setText(f"Word saved. PDF conversion failed: {e}")
            else:
                self.status_label.setText("Capture ended. No document to save.")
                logging.info("Capture ended. No document to save.")

            # Delete captured images if the checkbox is checked
            if self.delete_images_after_save:
                self.cleanup_captured_images()
        except Exception as e:
            logging.error(f"Error stopping capture: {e}")
            QMessageBox.critical(self, "Error", f"Error stopping capture: {e}")

    def convert_to_pdf(self):
        if self.doc_path and os.path.exists(self.doc_path):
            pdf_path = os.path.splitext(self.doc_path)[0] + ".pdf"
            try:
                convert(self.doc_path, pdf_path)
                self.status_label.setText(f"Word document converted to PDF: {pdf_path}")
                QMessageBox.information(self, "Conversion Successful",
                                        f"Word document successfully converted to PDF:\n{pdf_path}")
                logging.info(f"Word document converted to PDF: {pdf_path}")
            except Exception as e:
                logging.error(f"PDF conversion failed: {e}")
                self.status_label.setText(f"PDF conversion failed: {e}")
                QMessageBox.critical(self, "Conversion Failed",
                                     f"Error converting to PDF:\n{e}\n\nMake sure Microsoft Word is installed.")
        else:
            self.status_label.setText("No Word document has been saved yet.")
            QMessageBox.warning(self, "No Word Document", "Please capture screenshots and save the Word document first.")
            logging.warning("No Word document has been saved yet.")

    def cleanup_captured_images(self):
        for img_path in self.captured_images:
            try:
                os.remove(img_path)
                print(f"Deleted: {img_path}")
                logging.info(f"Deleted image: {img_path}")
            except Exception as e:
                print(f"Error deleting {img_path}: {e}")
                logging.error(f"Error deleting image: {img_path}: {e}")
        self.captured_images = []

    def generate_excel(self):
        if not self.excel_path:
            return

        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet.append(["Co.", "Description", "Image"])  # Header row

        for row_num, data in enumerate(self.captured_data):
            sheet.cell(row=row_num + 2, column=1, value=data["co"])
            sheet.cell(row=row_num + 2, column=2, value=data["description"])
            try:
                img = ExcelImage(data["image_path"])
                #  Anchoring the image to the cell.
                img.anchor = sheet.cell(row=row_num + 2, column=3).coordinate
                sheet.add_image(img)

                # Adjust row height to fit the image (optional)
                img_height_px = Image.open(data["image_path"]).height
                row_height_points = img_height_px * 72 / 96  # Approximate conversion
                sheet.row_dimensions[row_num + 2].height = row_height_points

                # Adjust column width for better image visibility (optional)
                sheet.column_dimensions['C'].width = 40  # Adjust as needed.
                logging.info(f"Added image to Excel: {data['image_path']}")
            except Exception as e:
                error_message = f"Error adding image to Excel: {e}"
                print(error_message)
                logging.error(error_message)
                sheet.cell(row=row_num + 2, column=3, value=f"Error: {e}")

        try:
            workbook.save(self.excel_path)
            self.status_label.setText(f"Excel document with images generated: {self.excel_path}")
            QMessageBox.information(self, "Excel Generated",
                                    f"Excel document with images successfully generated:\n{self.excel_path}")
            logging.info(f"Excel document with images generated: {self.excel_path}")
        except Exception as e:
            error_message = f"Error generating Excel document with images: {e}"
            self.status_label.setText(error_message)
            QMessageBox.critical(self, "Excel Generation Failed", error_message)
            logging.error(error_message)

    def capture_screenshot(self):
        if not self.capture_enabled or self.doc is None:
            return

        folder = self.folder_input.text().strip()
        if not os.path.exists(folder):
            os.makedirs(folder)

        case_name = self.test_case_input.text().strip() or "Evidence"
        description = self.description_input.text().strip()
        co = self.screenshot_count  # Capture current count for Excel
        if self.timestamp_checkbox.isChecked():
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            description += f"_{timestamp}"

        filename_base = f"{case_name}_{self.screenshot_count}_{description}".replace(" ", "_")
        screens = QApplication.screens()

        if self.capture_mode == "single":
            selected_index = self.single_monitor_combo.currentIndex()
            if 0 <= selected_index < len(screens):
                screen = screens[selected_index]
                geometry = screen.geometry()
                image_path = os.path.join(folder, f"{filename_base}_monitor_{selected_index + 1}.png")
                screenshot = pyautogui.screenshot(
                    region=(geometry.x(), geometry.y(), geometry.width(), geometry.height()))
                screenshot.save(image_path)
                self.add_to_word(image_path,
                                 f"Screenshot {self.screenshot_count} (Monitor {selected_index + 1}): {description}")
                self.captured_images.append(image_path)
                self.captured_data.append({"co": co, "description": description, "image_path": image_path})
                self.update_preview(image_path)
                logging.info(f"Single monitor screenshot captured: {image_path}")
        elif self.capture_mode == "all":
            all_x = min(s.geometry().x() for s in screens)
            all_y = min(s.geometry().y() for s in screens)
            all_width = max(s.geometry().width() + s.geometry().x() for s in screens) - all_x
            all_height = max(s.geometry().height() + s.geometry().y() for s in screens) - all_y
            image_path = os.path.join(folder, f"{filename_base}_all_monitors.png")
            screenshot = pyautogui.screenshot(region=(all_x, all_y, all_width, all_height))
            screenshot.save(image_path)
            self.add_to_word(image_path, f"Screenshot {self.screenshot_count} (All Monitors): {description}")
            self.captured_images.append(image_path)
            self.captured_data.append({"co": co, "description": description, "image_path": image_path})
            self.update_preview(image_path)
            logging.info(f"All monitors screenshot captured: {image_path}")
        elif self.capture_mode == "multiple":
            selected_indices = [i for i in range(self.multiple_monitor_list.count()) if
                                self.multiple_monitor_list.item(i).checkState() == Qt.Checked]
            for index in selected_indices:
                if 0 <= index < len(screens):
                    screen = screens[index]
                    geometry = screen.geometry()
                    image_path = os.path.join(folder, f"{filename_base}_monitor_{index + 1}.png")
                    screenshot = pyautogui.screenshot(
                        region=(geometry.x(), geometry.y(), geometry.width(), geometry.height()))
                    screenshot.save(image_path)
                    self.add_to_word(image_path,
                                     f"Screenshot {self.screenshot_count} (Monitor {index + 1}): {description}",
                                     new_page=True)
                    self.captured_images.append(image_path)
                    self.captured_data.append({"co": co, "description": description, "image_path": image_path})
                    self.update_preview(image_path)
                    logging.info(f"Multiple monitors screenshot captured: {image_path}")
                    if len(selected_indices) > 1 and index < selected_indices[-1]:
                        self.doc.add_page_break()
            if not selected_indices:
                self.status_label.setText("Please select at least one monitor in 'Select Multiple Monitors' mode.")
                logging.warning("Please select at least one monitor in 'Select Multiple Monitors' mode.")
                return

        self.status_label.setText(f"Screenshot {self.screenshot_count} captured and added to document.")
        if self.increment_checkbox.isChecked():
            self.screenshot_count += self.increment_spin.value()

    def add_to_word(self, image_path, description, new_page=False):
        if new_page and self.doc.paragraphs:
            self.doc.add_page_break()

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(description)
        run.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        try:
            # Calculate available width (adjust margins as needed)
            available_width = Inches(6.5)
            img = Image.open(image_path)
            img_width_px, img_height_px = img.size
            aspect_ratio = img_height_px / img_width_px
            scaled_width = available_width
            scaled_height = scaled_width * aspect_ratio
            self.doc.add_picture(image_path, width=scaled_width, height=scaled_height)
            last_paragraph = self.doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logging.info(f"Added image to Word document: {image_path}")
        except Exception as e:
            error_message = f"Error adding image to Word document: {e}"
            self.doc.add_paragraph(error_message)
            logging.error(error_message)

        self.doc.add_paragraph("")  # Add a blank line for spacing

    def update_preview(self, image_path):
        """Updates the preview label with the most recently captured screenshot."""
        pixmap = QPixmap(image_path)
        pixmap = pixmap.scaled(self.preview_label.size(), Qt.KeepAspectRatio, Qt.SmoothTransformation)
        self.preview_label.setPixmap(pixmap)
        logging.info(f"Updated preview with image: {image_path}")


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = ScreenshotApp()
    window.resize(640, 900)
    window.show()
    sys.exit(app.exec_())
